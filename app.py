import json
import requests
import os
import time
import cloudant
from cloudant import Cloudant
import pandas as pd
from docx import Document
from cloudant.error import CloudantException
from cloudant.result import Result, ResultByKey, QueryResult
from flask import Flask, request, make_response, render_template, session, g

user= "000d5b69-974a-4f7b-9118-5bcb7aed2484-bluemix"	
password= "ecab228d0b549390a55b708bcfe28f6818e223aa5761dc2628e65859c7211e87"	
host= "000d5b69-974a-4f7b-9118-5bcb7aed2484-bluemix.cloudant.com"
url = 'https://' + host
client = Cloudant(user, password, url=url, connect=True)
app=Flask(__name__)
  
@app.route('/docx')
def download_docx():
    with open("static/workinfo.docx",'rb') as f:
        body = f.read()
    response = make_response(body)
    response.headers["Content-Disposition"] = "attachment; filename= workinfo.docx"
    return response

@app.route('/webhook',methods=['POST'])
def webhook():
    req= request.get_json(silent=True, force=True)
    sessionId=req.get("sessionId")
    result=req.get("result")
    contexts=result.get("contexts")
    print(contexts)
    action=result.get("action")
    print(action)
    par=contexts[0].get("parameters")
    resolution = "Incomplete"
    if action == "troubleshooting.webhook" :
        resolution =  troubleshoot(par)
    if action == "healthcheck" :
        resolution = healthcheck(par)
    if action == "workinfo.creation" :
        resolution = workinfo(par)
    if action == "predictiveanalysis" :
        resolution = predictiveanalysis(par)
    if action == "updateproperties.beta" :
        resolution = "update props"
    
    res={"speech":resolution, "displayText":resolution, "source":"jarvis-chatbot"}
    res=json.dumps(res, indent=4)
    print(res)
    r = make_response(res)
    r.headers['Content-Type'] = 'application/json'
    return r

def troubleshoot(par):
    SERVER = par.get("SERVER")
    APPLICATION = par.get("APPLICATION")
    session = client.session()
    db = client['esb-data']
    query_ts = cloudant.query.Query(db,selector={"SERVER":SERVER,"APPLICATION":APPLICATION,"$or": [{"GET": "DISABLED"},{"PUT": "DISABLED"},{"CURDEPTH": {"$gt": 4000}},{"IP_PROCS": {"$gt": 0}},{"OP_PROCS": {"$gt": 0}},{"CH_STATUS": "STOPPED"},{"EG_STATUS": "RUNNING"},{"FLOW_STATUS": "RUNNING"},{"MSG_INSTANCES": {"$gt": 25}}]})
    time.sleep(1)
    queryresult_ts = QueryResult(query_ts)
    #print(queryresult_ts)
    for doc in queryresult_ts:
        try:
            issue="ISSUE WITH "+doc['FLOW_NAME']
            details=str(doc)
        except:
            issue="ISSUE WITH "+doc['QUEUE']
            details=str(doc)
    result = issue + details
    #details={}
    #log=client['jarvis-interaction']
    #doc=log.create_document(details)
    #doc.save()
    return result

def healthcheck(par):
    SERVER= par.get("SERVER")
    result = SERVER
    session=client.session()
    db=client['esb-data']
    query_hc = cloudant.query.Query(db,selector={"SERVER":SERVER,"CPU":{"$gt":0}})
    time.sleep(1)
    queryresult_hc = QueryResult(query_hc)
    for doc in queryresult_hc:
        print(doc["SERVER"]+" has CPU:"+str(doc["CPU"])+" and Memory:"+str(doc["MEMORY"]))
        print(doc["QMGR"]+" is "+doc["QM_STATUS"]+" and" + doc["CLUSTER_STATUS"] + " in the Cluster")
        print("Listener port :"+str(doc["LISTENER"])+" is "+doc["LISTENER_STATUS"])
        result_tmp = doc["SERVER"]+" is " + doc["CLUSTER_STATUS"] + " in the Cluster<br>"+"with CPU:"+str(doc["CPU"])+" and Memory:"+str(doc["MEMORY"]) + "<br>Broker/QMGR: "+doc["QMGR"]+" is "+doc["QM_STATUS"] + " with Listener port :"+str(doc["LISTENER"])+" is "+doc["LISTENER_STATUS"]
    return result_tmp

def workinfo(par):
    SERVER= par.get("SERVER")
    TASK= par.get("WORKINFO-TASK")
    CRQ=str(par.get("CRQ"))
    generate_docx(SERVER,TASK,CRQ)
    link="Please download below<br><a href=\"https://jarvis-walgreens-webhook.herokuapp.com/docx\" class=\"btn btn-primary\">DOWNLOAD</a>"
    return link
    
def predictiveanalysis(par):
    session=client.session()
    db=client['esb-data']
    query_pa_o1 = cloudant.query.Query(db,selector={"CPU":{"$gt":0}})
    health_report(query_pa_o1)
    time.sleep(1)
    for doc in query_pa_o1:
      SERVER = doc["SERVER"]
      print(SERVER)
      query_pa_o2 = cloudant.query.Query(db,selector={"SERVER":SERVER})
      time.sleep(1)
      for doc in query_pa_i1:
        print(doc["CPU"])
        print(doc["Memory"])
                
    return "From my analysis there seems to be no issues as of now"

def generate_docx(SERVER,TASK,CRQ):
    document = Document("static/template.docx")
    document.add_heading("CRQ Details:")
    document.add_paragraph("CRQ Number: "+CRQ)
    document.add_paragraph("Implementation Team: ESB 2L")
    document.add_heading("Task Details:")
    document.add_paragraph("Task:"+TASK)
    document.add_paragraph("SERVER:"+SERVER)
    
    if(TASK=="VALIDATION"):
        document.add_heading("Validation Plan:")
        document.add_paragraph("1.Check MQ/MB is running or not", style = 'ListNumber')
        document.add_paragraph("2.Check Listener port and its status", style = 'ListNumber')
        document.add_paragraph("3.Check all MQ/MB processes are running", style = 'ListNumber')
        document.add_paragraph("4.Check Tivoli agents are up and running", style = 'ListNumber')
    if(TASK=="BOUNCEMQ"):
        document.add_heading("Implementation Plan:")
        document.add_paragraph("1.Stop MQ services", style = 'ListNumber')
        document.add_paragraph("2.Check all MQ processes are down", style = 'ListNumber')
        document.add_paragraph("", style = 'ListNumber')
        document.add_paragraph("3.Start MQ services", style = 'ListNumber')
        document.add_paragraph("4.Check all MQ processes are UP and RUNNING", style = 'ListNumber')
    document.add_paragraph("Step 2:", style = 'ListNumber')
    if(TASK=="BOUNCEMB"):
        document.add_heading("Implementation Plan:")
        document.add_paragraph("1.Stop MB services", style = 'ListNumber')
        document.add_paragraph("2.Check all MB processes are down", style = 'ListNumber')
        document.add_paragraph("", style = 'ListNumber')
        document.add_paragraph("3.Start MB services", style = 'ListNumber')
        document.add_paragraph("4.Check all MB processes are UP and RUNNING", style = 'ListNumber')
    document.add_paragraph("Close the task in Remedy", style = 'ListNumber')
    document.save("static/workinfo.docx")

def health_report(QUERY_RES):
    document = Document("static/template.docx")
    document.add_heading("Health check Report:")
    table = document.add_table(rows=1, cols=8, style = 'TableGrid')
    hdr_cells[0].text = 'SERVER'
    hdr_cells[1].text = 'CPU'
    hdr_cells[2].text = 'MEMORY'
    hdr_cells[3].text = 'BROKER/QMGR'
    hdr_cells[4].text = 'BROKER/QMGR STATUS'
    hdr_cells[5].text = 'CLUSTER STATUS'
    hdr_cells[6].text = 'LSR PORT'
    hdr_cells[7].text = 'LSR STATUS'
    for doc in QUERY_RES:
        row_cells = table.add_row().cells
        row_cells[0].text = doc['SERVER']
        row_cells[1].text = doc['CPU']
        row_cells[2].text = doc['MEMORY']
        row_cells[3].text = doc['QMGR']
        row_cells[4].text = doc['QM_STATUS']
        row_cells[5].text = doc['CLUSTER_STATUS']
        row_cells[6].text = doc['LISTENER']
        row_cells[7].text = doc['LISTENER_STATUS']
    document.save("static/HEALTHCHECK.docx")

port = os.getenv('VCAP_APP_PORT', '5000')
if __name__ == "__main__":
    app.run(host='0.0.0.0', port=int(port), use_reloader=True, debug=True)
